"""Build the Phase 2 submission .docx in the same style as the example
(Sathyamurthi Karthikeyan's submission). Run via:

    uv run --with python-docx python docs/build_submission_docx.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches


STUDENT = "Ashok Kumar Meena"
TITLE = "German Document Type Classification using CNNs"
PROGRAM = "M.Sc. Software Engineering"
INSTITUTION = "University of Europe for Applied Sciences, Potsdam"
SUPERVISOR = "Raja Hashim Ali"
DATE = "08.06.2026"
GITHUB_URL = "https://github.com/ashok-cse/PR_SE26_German_Document_Type_Classification_with_CNNs"
KAGGLE_URL = "https://www.kaggle.com/code/ashokkrcse/german-document-type-classification-with-cnns?scriptVersionId=325254710"

OUT_PATH = Path(__file__).resolve().parent / (
    f"Phase 2 - Proposal and Code Implementation - {STUDENT}.docx"
)

doc = Document()

# Default body font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def add_heading(text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"


def add_para(text: str, *, bold: bool = False, align=None, size: int | None = None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)


# ---------------------------------------------------------------------------
# Title page
# ---------------------------------------------------------------------------

for _ in range(2):
    doc.add_paragraph()

add_para(TITLE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=20)
for _ in range(3):
    doc.add_paragraph()

add_para(f"Submitted By: {STUDENT}", align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
for _ in range(3):
    doc.add_paragraph()

add_para(PROGRAM, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para(INSTITUTION, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
for _ in range(3):
    doc.add_paragraph()

add_para(f"Supervisor: {SUPERVISOR}", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para(f"Date: {DATE}", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

doc.add_page_break()

# ---------------------------------------------------------------------------
# Contents
# ---------------------------------------------------------------------------

add_heading("Contents", level=1)
toc_entries = [
    ("Background and Motivation", 3),
    ("Problem Statement", 3),
    ("Dataset Used", 4),
    ("Research Questions", 4),
    ("Proposed Methodology", 5),
    ("CNN Model Design", 5),
    ("Transfer Learning Model", 5),
    ("Evaluation Metrics", 6),
    ("Expected Results", 6),
]
for name, page in toc_entries:
    p = doc.add_paragraph()
    run = p.add_run(f"{name}\t{page}")
    run.font.size = Pt(11)

doc.add_page_break()

# ---------------------------------------------------------------------------
# Background and Motivation
# ---------------------------------------------------------------------------

add_heading("Background and Motivation", level=1)
doc.add_paragraph(
    "Document classification plays a vital role in digitisation pipelines of "
    "administrative, legal, scientific, and financial archives. German-speaking "
    "institutions such as the Deutsche Digitale Bibliothek, the Bundesarchiv, "
    "EU tender portals, and German regulatory bodies process millions of "
    "scanned pages every year. Once a document is scanned, the very first "
    "downstream decision is almost always — what kind of document is this? "
    "Routing an invoice differs from routing a patent, which again differs "
    "from routing a regulatory notice. Doing this routing manually is slow, "
    "expensive, and error-prone, and is therefore an ideal task to automate "
    "with image-based machine learning."
)
doc.add_paragraph(
    "Most published work on document image classification has focused on "
    "English-language benchmarks such as RVL-CDIP and Tobacco-3482. German "
    "documents pose a visually distinct problem: longer compound words alter "
    "line-wrapping patterns, German official documents use different masthead "
    "and header conventions, and legal Verordnungen carry different paragraph "
    "densities than scientific Fachartikel. A model trained purely on English "
    "documents may not transfer cleanly to these layouts. The previous "
    "biometric- and OCR-based pipelines for archive routing relied on "
    "text-only features and were sensitive to scan quality. Image-based "
    "classification using Convolutional Neural Networks (CNNs) solves this "
    "by reading the visual layout — masthead position, column structure, "
    "table density, figure placement — directly from the page, in a way "
    "that does not depend on perfect OCR."
)
doc.add_paragraph(
    "The most recent advancements include applying transfer learning "
    "(ResNet, EfficientNet, ViT) to document images, which gives better "
    "results with much less training data. This project applies these "
    "concepts specifically to the six-class DocLayNet-base dataset to build and "
    "rigorously evaluate a CNN-based classifier for German scanned documents."
)

# ---------------------------------------------------------------------------
# Problem Statement
# ---------------------------------------------------------------------------

add_heading("Problem Statement", level=1)
doc.add_paragraph(
    "German document classification is one of the most important and growing "
    "tasks for automating archive workflows in German-speaking institutions. "
    "Traditional approaches using OCR plus rule-based text classifiers find "
    "it difficult to cope with poor scan quality, mixed German fonts, and "
    "the strong layout variability across categories such as financial "
    "reports, manuals, patents, and government tenders. CNN-based "
    "classification solves this problem by extracting spatial features "
    "directly from the page image: the convolution layers learn local "
    "layout features, pooling layers reduce dimensionality, the fully "
    "connected layers combine the features, and the softmax layer "
    "produces a probability distribution over the six document types."
)
doc.add_paragraph(
    "Given a single-page scanned German document image x, the system "
    "predicts its document type y from six DocLayNet categories: "
    "financial_reports, scientific_articles, laws_and_regulations, "
    "government_tenders, manuals, and patents. Only visual layout features "
    "are used — no OCR text features are introduced. We train a custom CNN "
    "from scratch as a baseline and compare it against a ResNet50 transfer "
    "learning model (ImageNet pre-trained) to measure the gap and the "
    "robustness of each approach to phone-camera-style image degradations."
)

# ---------------------------------------------------------------------------
# Dataset Used
# ---------------------------------------------------------------------------

add_heading("Dataset Used", level=1)
doc.add_paragraph(
    "The dataset used is DocLayNet-base, the Parquet-format mid-size variant "
    "of DocLayNet released by IBM Research under the CDLA-Permissive-1.0 "
    "license. DocLayNet-base provides 8,057 scanned pages at 1025×1025 px "
    "standardised resolution (train 6,910; validation 648; test 499) and is "
    "small enough to download in full on a Kaggle GPU notebook while still "
    "covering all six functional document categories. Each page is labelled "
    "with one of six document-source categories that map directly to "
    "functional document types:"
)
classes = [
    "financial_reports (annual reports, SEC filings)",
    "scientific_articles (research papers)",
    "laws_and_regulations (statutes, FAA regulations, compliance documents)",
    "government_tenders (EU TED tender documents)",
    "manuals (technical documentation, user guides)",
    "patents (granted patents)",
]
for c in classes:
    doc.add_paragraph(c, style="List Bullet")

doc.add_paragraph(
    "Link – DocLayNet-base (HuggingFace): "
    "https://huggingface.co/datasets/pierreguillou/DocLayNet-base"
)

add_heading("German-Language Filtering", level=2)
doc.add_paragraph(
    "DocLayNet does not expose an explicit language field, and the corpus is "
    "heavily English-dominant: only about 2.5 % of pages across the full "
    "DocLayNet are German (the remainder are mostly English, with some French "
    "and Japanese). Because this project targets German documents, the notebook "
    "applies a strict German-language filter rather than training on the full "
    "multilingual corpus. For each page, the annotation text is concatenated and "
    "passed through the langdetect library; a page is kept only when it is "
    "confidently classified as German — detection probability p(de) ≥ 0.90 with "
    "at least 40 characters of text (controlled by the STRICT_GERMAN, "
    "GERMAN_MIN_PROB and GERMAN_MIN_TEXT settings). The language detector is "
    "seeded for reproducibility."
)
doc.add_paragraph(
    "This keeps the “German” claim accurate at the cost of a smaller training "
    "set. The notebook emits a per-class German-yield report — a table of pages "
    "scanned (all languages) versus German pages kept, with the German rate per "
    "class (figures/t0_german_yield.csv), and a companion grouped-bar figure "
    "(figures/f0_german_yield.png). This makes the size of the genuinely-German "
    "subset explicit and surfaces any class that becomes under-represented after "
    "filtering. Note that visual layout is largely language-independent, so the "
    "CNN methodology is unaffected; the filter changes which pages train the "
    "model, not how it learns. The 8,057-page distribution above describes the "
    "pre-filter DocLayNet-base corpus."
)

add_heading("Implementation Note on Dataset Loading", level=2)
doc.add_paragraph(
    "The original Phase 2 proposal targeted the ≈ 2,000 German pages within "
    "DocLayNet-large (≈ 80k pages). Two upstream bugs in the public DocLayNet "
    "HuggingFace distributions forced a two-step pivot during implementation:"
)
doc.add_paragraph(
    "(1) DocLayNet-large is distributed only via a custom Python loading "
    "script. In streaming mode that script raises FileNotFoundError on "
    "remote zip://...::https://... URLs (it expects the zip to be already "
    "extracted locally), and the script-based loader format itself is "
    "rejected by the current `datasets` ≥ 4.0 library."
)
doc.add_paragraph(
    "(2) The mid-size variant DocLayNet-base also ships a custom loader "
    "script, which fails during Arrow conversion with "
    "`ArrowInvalid: Float value … was truncated converting to int64` — its "
    "declared schema casts bounding-box coordinates to int64 while the "
    "underlying data contains floats."
)
doc.add_paragraph(
    "The notebook therefore bypasses `datasets.load_dataset` entirely. It "
    "uses `huggingface_hub.hf_hub_download` to pull the underlying "
    "`dataset_base.zip` (~3.8 GB), extracts it into Kaggle working storage, "
    "walks the JSON annotation files for the `doc_category` label, locates "
    "the matching PNG, and persists each page as a 224×224 JPEG. The CNN "
    "architecture, transfer-learning protocol, evaluation metrics, and "
    "Grad-CAM analysis are unchanged from the proposal; only the language "
    "filter is dropped (no public language-tagged DocLayNet variant is "
    "currently maintained). The motivation for German document routing "
    "remains the principal downstream use case for the trained classifier."
)

# ---------------------------------------------------------------------------
# Research Questions
# ---------------------------------------------------------------------------

add_heading("Research Questions", level=1)
rqs = [
    "How accurately can CNN-based models classify German scanned documents "
    "into the six DocLayNet categories using only visual layout features?",
    "Which model family — a custom CNN trained from scratch on the small "
    "DocLayNet-base dataset, or a ResNet50 fine-tuned from ImageNet — gives the best "
    "balance between predictive performance and computational efficiency?",
    "How do preprocessing, data augmentation, and class-imbalance handling "
    "influence overall accuracy and per-class robustness?",
    "Do Grad-CAM explanations show that the trained models attend to "
    "meaningful page regions (mastheads, headers, table boundaries) when "
    "predicting the document type?",
    "How robust is the best model to phone-camera-style degradations "
    "(perspective warp, JPEG compression, Gaussian blur, brightness shift) "
    "that simulate the visual gap between flatbed scans and mobile-phone "
    "captures of paper documents?",
]
for rq in rqs:
    doc.add_paragraph(rq, style="List Bullet")

# ---------------------------------------------------------------------------
# Proposed Methodology
# ---------------------------------------------------------------------------

add_heading("Proposed Methodology", level=1)
doc.add_paragraph(
    "The methodology compares a custom CNN against a ResNet50 transfer "
    "learning model on DocLayNet-base to test both predictive performance "
    "and robustness. The dataset (≈ 8,000 pages, six document types) is "
    "downloaded from HuggingFace as a Parquet dataset and persisted as "
    "224×224 JPEGs on local Kaggle storage so the downstream tf.data "
    "pipelines can iterate the files lazily. Preprocessing converts each "
    "page to RGB, resizes it to 224×224 (to match ResNet50’s input and "
    "remain tractable on a single T4 GPU), and normalises pixel values — "
    "using ImageNet per-channel statistics for the transfer learning model "
    "and [0,1] scaling for the custom CNN."
)
doc.add_paragraph(
    "Class labels are integer-encoded in the canonical order "
    "(financial_reports = 0 … patents = 5). The split is stratified: 70 % "
    "training, 15 % validation, 15 % test. To address the natural class "
    "imbalance, class-weighted categorical cross-entropy is used "
    "(weight = total / (n_classes × class_count)). Training-time "
    "augmentation reduces overfitting through small random rotations, "
    "translations, zooms, and brightness shifts."
)
doc.add_paragraph(
    "The transfer learning model uses ResNet50 with ImageNet pre-trained "
    "weights for faster convergence and better feature extraction. The "
    "model is trained with the Adam optimiser, a batch size of 32, and "
    "early stopping monitored on validation accuracy. ResNet50 is trained "
    "in two stages: first the classifier head is trained while the "
    "backbone is frozen (10 epochs, lr = 1e-3), and then the last residual "
    "block (conv5_*) is unfrozen and fine-tuned (20 epochs, lr = 1e-5). "
    "Evaluation reports precision, recall, F1-score, overall accuracy, ROC "
    "curves with macro and per-class AUC, the confusion matrix to inspect "
    "misclassification patterns, and Grad-CAM heatmaps to visualise which "
    "page regions the model attended to."
)

# ---------------------------------------------------------------------------
# CNN Model Design
# ---------------------------------------------------------------------------

add_heading("CNN Model Design", level=1)
doc.add_paragraph(
    "The custom CNN is designed as a deliberately simple baseline so the "
    "value of transfer learning can be quantified. The input layer accepts "
    "224×224×3 RGB images. The convolutional body is organised as four "
    "blocks with 32, 64, 128, and 256 filters respectively; each block "
    "uses a 3×3 convolution, batch normalisation, ReLU activation, and "
    "2×2 max-pooling. Global average pooling is then used to flatten the "
    "feature map into a vector (which keeps the parameter count low and "
    "reduces overfitting risk on this small dataset). A fully connected "
    "dense layer with 128 units and ReLU activation is followed by a "
    "Dropout(0.5) layer to reduce overfitting, and finally a Dense(6) "
    "layer with softmax activation produces the per-class probability "
    "distribution. The total parameter count is approximately 1 million, "
    "and the model is trained from scratch with the Adam optimiser at a "
    "learning rate of 1e-3."
)

# ---------------------------------------------------------------------------
# Transfer Learning Model
# ---------------------------------------------------------------------------

add_heading("Transfer Learning Model", level=1)
doc.add_paragraph(
    "ResNet50 is used as the transfer learning model for classifying the "
    "German document images. ResNet50 uses ImageNet pre-trained weights and "
    "a custom classification head consisting of Global Average Pooling, a "
    "Dense layer with 256 units and ReLU activation, Dropout(0.4) to "
    "reduce overfitting, and a final Dense(6) layer with softmax output. "
    "Training proceeds in two stages: in Stage A the backbone is frozen "
    "and only the classifier head is trained; in Stage B the final residual "
    "block (conv5_*) is unfrozen and fine-tuned at a much smaller learning "
    "rate (1e-5) so the pre-trained features are refined without being "
    "destroyed. This two-stage strategy is well-suited to a small target "
    "dataset like DocLayNet-base and reduces the risk of catastrophic "
    "forgetting."
)

# ---------------------------------------------------------------------------
# Evaluation Metrics
# ---------------------------------------------------------------------------

add_heading("Evaluation Metrics", level=1)
doc.add_paragraph(
    "The models are evaluated using overall accuracy, precision, recall, "
    "and F1-score (macro-averaged because of the class imbalance) to "
    "measure label correctness, model sensitivity, and the harmonic mean "
    "of precision and recall, respectively. A confusion matrix (both "
    "absolute counts and row-normalised) is used to analyse misclass-"
    "ification patterns across the six classes. ROC curves are plotted "
    "in a one-vs-rest formulation for multi-class evaluation; the True "
    "Positive Rate vs. False Positive Rate curve is shown per class, and "
    "the area under the curve (AUC) is reported per class and macro-"
    "averaged. Model efficiency is measured using the total number of "
    "trainable parameters and the per-step training time. Robustness is "
    "additionally measured on a phone-camera-simulated test set, and the "
    "drop in macro-F1 between clean and simulated images is reported."
)

# ---------------------------------------------------------------------------
# Results
# ---------------------------------------------------------------------------

add_heading("Results", level=1)
doc.add_paragraph(
    "Both models are evaluated on the local DocLayNet-base page-image dataset "
    "(8,057 images, split 70/15/15). The reported values below are generated "
    "from the saved models on the stratified test split of 1,209 images."
)
doc.add_paragraph(
    "The Custom CNN is the smallest model (423,526 parameters), but the saved "
    "model performs poorly on the test set: accuracy 0.326, macro-F1 0.082, "
    "and macro-AUC 0.479. Its confusion matrix shows collapse toward the "
    "financial_reports class."
)
doc.add_paragraph(
    "The ResNet50 transfer-learning model is substantially stronger: accuracy "
    "0.761, macro-F1 0.724, and macro-AUC 0.956 on the same test split. It has "
    "24,113,798 parameters and is the primary model for inference."
)
doc.add_paragraph(
    "Under the phone-camera-style robustness evaluation, ResNet50 macro-F1 "
    "drops from 0.724 to 0.606. The Custom CNN remains at 0.082 macro-F1 "
    "because its clean-set predictions are already degenerate."
)

# Figures — embed each PNG followed by its centred caption.
add_heading("Figures", level=1)
doc.add_paragraph(
    "The figures below are actual artifacts generated from the local dataset "
    "and saved models by scripts/evaluate_actual_figures.py."
)

FIG_DIR = Path(__file__).resolve().parent.parent / "figures"
figure_specs = [
    ("f1_class_distribution.png", "Fig 1. DocLayNet-base class distribution."),
    ("f2_sample_per_class.png", "Fig 2. Sample page per class."),
    ("f3_custom_cnn_curves.png",
     "Fig 3. Custom CNN training and validation accuracy and loss curves."),
    ("f4_resnet50_curves.png",
     "Fig 4. ResNet50 training and validation accuracy and loss curves "
     "(Stages A and B)."),
    ("f5_custom_cnn_confusion.png",
     "Fig 5. Custom CNN confusion matrix (counts and row-normalised)."),
    ("f6_custom_cnn_roc.png",
     "Fig 6. Custom CNN ROC curves (False Positive vs True Positive) with "
     "per-class AUC."),
    ("f7_resnet50_confusion.png",
     "Fig 7. ResNet50 confusion matrix (counts and row-normalised)."),
    ("f8_resnet50_roc.png",
     "Fig 8. ResNet50 ROC curves (False Positive vs True Positive) with "
     "per-class AUC."),
    ("f9_model_comparison.png",
     "Fig 9. Model comparison: accuracy, macro-F1, and macro-AUC for "
     "Custom CNN vs. ResNet50."),
    ("f10_gradcam.png",
     "Fig 10. Grad-CAM overlays for ResNet50, one example page per class."),
    ("f11_robustness.png",
     "Fig 11. Robustness comparison: clean macro-F1 vs. phone-camera-"
     "simulated macro-F1."),
]

for filename, caption in figure_specs:
    img_path = FIG_DIR / filename
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(img_path), width=Inches(6.0))
    cap_p = doc.add_paragraph(caption)
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap_p.runs:
        run.italic = True
    doc.add_paragraph()  # spacing between figures

# ---------------------------------------------------------------------------
# Links
# ---------------------------------------------------------------------------

add_heading("Repository and Notebook Links", level=1)
doc.add_paragraph(f"GitHub – {GITHUB_URL}")
doc.add_paragraph(f"Kaggle – {KAGGLE_URL}")


# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------

doc.save(OUT_PATH)
print(f"Wrote {OUT_PATH}")
